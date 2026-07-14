from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "development-documentation.docx"
RED = RGBColor(198, 29, 35)
BLACK = RGBColor(20, 20, 20)
GRAY = RGBColor(91, 91, 91)
LIGHT = "F4F1EE"
WIDTHS = [Inches(1.45), Inches(5.05)]


def set_font(run, size, color=BLACK, bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table):
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, WIDTHS[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side in ("top", "start", "bottom", "end"):
                margin = OxmlElement(f"w:{side}")
                margin.set(qn("w:w"), "110")
                margin.set(qn("w:type"), "dxa")
                margins.append(margin)
            tc_pr.append(margins)


def add_paragraph(doc, text="", size=11, color=BLACK, bold=False, italic=False, after=8, before=0, alignment=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = 1.2
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    set_font(run, size, color, bold, italic)
    return paragraph


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 13, 3: 11.5}
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, sizes[level], RED if level < 3 else BLACK, True)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, 10.5, BLACK)
    return paragraph


def add_label_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        shade(cells[0], LIGHT)
        label_run = cells[0].paragraphs[0].add_run(label)
        value_run = cells[1].paragraphs[0].add_run(value)
        set_font(label_run, 10, BLACK, True)
        set_font(value_run, 10, BLACK)
    set_table_geometry(table)
    return table


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DISRUPTIVE-STYLE NEXT.JS EXPERIENCE")
    set_font(run, 8, GRAY, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Coursework development documentation")
    set_font(run, 8, GRAY)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_header_footer(section)

    title = add_paragraph(doc, "Disruptive-Style Next.js Experience", 27, BLACK, True, after=6)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_paragraph(doc, "Development Documentation", 14, RED, True, after=18)
    add_label_table(doc, [
        ("Project Type", "Coursework web experience recreation"),
        ("Framework", "Next.js 14 with the App Router"),
        ("Deliverable", "Responsive homepage, lead route, and project documentation"),
        ("Prepared", "July 2026")
    ])
    add_paragraph(doc, "Purpose", 10, RED, True, after=4, before=22)
    add_paragraph(doc, "This document records the architecture, experience decisions, interaction design, and verification approach for the revised Next.js project. It is intended as a concise professional handoff for coursework review.", 11, BLACK, after=8)

    add_heading(doc, "1. Project Overview")
    add_paragraph(doc, "The implementation recreates the observed structure and feel of a performance marketing agency homepage using independently authored React, TypeScript, CSS, and animation. The goal is a high-fidelity learning exercise with a responsive, production-style code structure.")
    add_heading(doc, "Delivered Scope", 2)
    for text in [
        "Responsive Next.js homepage with reusable content data and client-side interaction state.",
        "Fixed white header, red-and-black wordmark treatment, desktop mega menus, and mobile navigation drawer.",
        "Animated dark showreel built with CSS shapes and keyframes to preserve the motion-first opening experience.",
        "Lead capture form connected to a Next.js route handler at /api/audit.",
        "Client marquee, reviews, proof points, service positioning, timeline, form, and footer sections."
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2. Technical Architecture")
    add_label_table(doc, [
        ("Presentation", "app/page.tsx provides the homepage sections, responsive navigation state, and audit form behavior."),
        ("Styling", "app/globals.css owns the visual tokens, responsive layouts, desktop/mobile behavior, marquee, and showreel animation."),
        ("Backend", "app/api/audit/route.ts accepts the audit request and returns a validation-aware confirmation response."),
        ("Metadata", "app/layout.tsx imports global styling and defines the application metadata."),
        ("Dependencies", "Next.js, React, TypeScript, and Lucide React are installed through the project package manifest.")
    ])

    add_heading(doc, "3. Visual and Interaction Decisions")
    add_heading(doc, "Reference-led visual system", 2)
    add_paragraph(doc, "A live-site audit informed the revised structure: a fixed white header, a red wordmark and call-to-action, uppercase navigation, a dark animated opening showreel, a brand marquee, proof-oriented editorial sections, and a dark footer. The first implementation's teal card-oriented system was replaced to align with these observed patterns.")
    add_heading(doc, "Navigation", 2)
    for text in [
        "Desktop menus toggle one full-width mega panel beneath the header and use grouped link columns.",
        "The mobile breakpoint replaces the horizontal menu with a right-side drawer, a backdrop, expandable groups, and a dedicated close control.",
        "In-page calls to action close active navigation state before routing to the audit section."
    ]:
        add_bullet(doc, text)
    add_heading(doc, "Motion", 2)
    add_paragraph(doc, "The opening showreel is an original CSS composition. Animated typography, an illustrated central figure, shifting numerical metrics, a grid treatment, and intermittent brand labels create continuous motion without copying source video media. It also avoids a static hero image, which was a primary concern in the revision request.")

    add_heading(doc, "4. Lead Capture Flow")
    add_paragraph(doc, "The audit form captures first name, business email, company, and annual revenue. The client creates a JSON payload and posts it to the local route handler. The route validates the request and sends a human-readable response that is displayed beneath the submit button.")
    add_paragraph(doc, "Extension point: the route handler is the correct place to connect a CRM, transactional email platform, data store, rate limiter, and server-side spam protection before production deployment.", 10.5, GRAY, italic=True)

    add_heading(doc, "5. Local Run and Verification")
    add_heading(doc, "Commands", 2)
    add_paragraph(doc, "npm.cmd install", 10.5, BLACK, bold=True, after=2)
    add_paragraph(doc, "npm.cmd run dev", 10.5, BLACK, bold=True, after=2)
    add_paragraph(doc, "npm.cmd run build", 10.5, BLACK, bold=True, after=8)
    add_heading(doc, "Verification Checklist", 2)
    for text in [
        "Open http://localhost:3000 and confirm the showreel animates on initial load.",
        "Open each desktop navigation group and confirm the mega panel appears below the fixed header.",
        "Reduce the viewport and confirm the off-canvas drawer opens, expands groups, and closes with its backdrop or close button.",
        "Submit the audit form with valid values and confirm the route response appears in the form status region.",
        "Run the production build to validate TypeScript compilation and route generation."
    ]:
        add_bullet(doc, text)

    add_heading(doc, "6. Scope and Asset Notice")
    add_paragraph(doc, "This repository is an educational recreation. It does not include copied website source code, copied video, or source-owned brand assets from the reference site. The visual motion, wordmark treatment, illustrations, and implementation are independently authored for coursework demonstration.")
    doc.core_properties.title = "Disruptive-Style Next.js Experience - Development Documentation"
    doc.core_properties.author = "Coursework Project"
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
